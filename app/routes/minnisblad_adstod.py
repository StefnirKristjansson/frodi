"""This module contains the routes for the minnisblad application."""

import os
from datetime import datetime
import json
from tempfile import NamedTemporaryFile
from fastapi import (
    APIRouter,
    Request,
    File,
    UploadFile,
    HTTPException,
    Form,
    Depends,
    status,
)
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.templating import Jinja2Templates
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI


templates = Jinja2Templates(directory="app/templates")

router = APIRouter()

load_dotenv()
BEARER_TOKEN = os.getenv("BEARER_TOKEN")

# Security scheme and token validation
token_auth_scheme = HTTPBearer()


def get_token(credentials: HTTPAuthorizationCredentials = Depends(token_auth_scheme)):
    """Validate the token and return it if it is valid."""
    if credentials.scheme != "Bearer":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication scheme.",
        )
    if credentials.credentials != BEARER_TOKEN:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid or expired token.",
        )
    return credentials.credentials


openai_api_key = os.getenv("OPENAI_API_KEY")

# Set the OpenAI API key globally

OpenAI.api_key = openai_api_key
client = OpenAI()


@router.get("/minnisblad-adstod", response_class=HTMLResponse)
async def minnisblad_adstod(request: Request):
    """Render the minnisblad-adstod page."""
    return templates.TemplateResponse("minnisblad-adstod.html", {"request": request})


@router.post("/minnisblad-adstod/upload/")
async def upload_file(
    request: Request,
    file: UploadFile = File(...),
    chapters: str = Form(...),
    _: str = Depends(get_token),
):
    """Process the uploaded file and return the modified document."""
    if (
        file.content_type
        != "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or not file.filename.endswith(".docx")
    ):
        raise HTTPException(
            status_code=400, detail="Invalid file type. Please upload a .docx file."
        )

    # Ensure the file pointer is at the beginning
    file.file.seek(0)
    text = extract_text_from_docx(file.file)

    word_count = len(text.split())
    if word_count > 5000 or word_count < 10:
        raise HTTPException(
            status_code=400,
            detail="The document must contain between 10 and 5000 words.",
        )

    try:
        selected_chapters = json.loads(chapters)
        respond_format = create_response_format(selected_chapters)
        openai_response = await send_text_to_openai(text, respond_format)
        # the openai_response is a string, so we need to convert it to a dictionary

        output_file_path = create_docx_from_json(openai_response)
        # create a timestamp in human readable format
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = "Frodi_minnisblad_" + timestamp + ".docx"

        return FileResponse(
            output_file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
        )
    except Exception as e:  # pylint: disable=broad-except
        return templates.TemplateResponse(
            "index.html", {"request": request, "error": str(e)}
        )


def extract_text_from_docx(file):
    """Extract text from a .docx file."""
    doc = Document(file)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)


async def send_text_to_openai(
    text: str, response_format: dict
) -> dict:  # pragma: no cover
    """Send the text to the OpenAI API and return the response."""
    # Using the OpenAI client as per your original
    content_text = """Notendinn sendi þér minnisblað, farðu mjög varlega yfir það og 
    finndu dæmi um önnur minnisblöð, 
    gefðu þér tíma að skoa Íslenskt málfar og stafsetningu."""
    messages = [
        {
            "role": "system",
            "content": [
                {
                    "type": "text",
                    "text": content_text,
                }
            ],
        },
        {
            "role": "user",
            "content": [{"type": "text", "text": text}],
        },
    ]

    completion = client.chat.completions.create(
        model="gpt-4o",
        messages=messages,
        temperature=1,
        max_tokens=2048,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
        response_format=response_format,
    )
    response = json.loads(completion.choices[0].message.content)
    return response


def create_response_format(selected_chapters: list) -> dict:
    """Create the response format based on the selected chapters."""
    base_response = base_response = {
        "type": "json_schema",
        "json_schema": {
            "name": "Minnisblad_adstod",
            "schema": {  # Correcting to include 'schema' directly
                "title": "Minnisblad_adstod",
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "The name of the schema"},
                    "type": {
                        "type": "string",
                        "enum": [
                            "object",
                            "array",
                            "string",
                            "number",
                            "boolean",
                            "null",
                        ],
                        "description": "The type of the schema (should match JSON Schema types)",
                    },
                    "properties": {
                        "type": "object",
                        "properties": {
                            "malfar": {
                                "type": "string",
                                "description": "Hérna setur þú inn hvað meigi bæta og breyta varðandi málfar, stíl og uppbyggingu minnisblaðsins.",
                            },
                            "stafsetning": {
                                "type": "string",
                                "description": "Hérna setur þú inn hvað meigi bæta og breyta varðandi stafsetningu passaðu að benda skýrt á allar stafsetningarvillur með að vitna í textan sem á við.",
                            },
                            "radleggingar": {
                                "type": "string",
                                "description": "Hérna setur þú inn hvað meigi bæta og breyta varðandi almennar raðleggingar. Farðu vel yfir hvernig minnisblöð eru yfir höfuð og hafðu það í huga",
                            },
                        },
                        "required": ["malfar", "stafsetning", "radleggingar"],
                        "additionalProperties": False,
                    },
                },
                "required": ["name", "type", "properties"],
                "additionalProperties": False,
            },
        },
    }

    return base_response


def create_docx_from_json(response_json: dict) -> str:
    """Create a Word document from the JSON response."""
    # Create a new Word document
    doc = Document()

    # Add the title
    doc.add_heading(response_json.get("titill", "Titill Ekki Tiltækur"), level=1)

    # if there is an introduction, add it to the document
    if "Inngangur" in response_json:
        doc.add_heading("Inngangur", level=2)
        doc.add_paragraph(response_json["Inngangur"])

    # if there is markmid, add it to the document
    if "Markmið" in response_json:
        doc.add_heading("Markmið", level=2)
        doc.add_paragraph(response_json["Markmið"])

    if "Áætlun" in response_json:
        doc.add_heading("Áætlun", level=2)
        doc.add_paragraph(response_json["Áætlun"])

    # if there are chapters, add them to the document
    if "kaflar" in response_json:
        for chapter in response_json["kaflar"]:
            doc.add_heading(chapter["chapter_title"], level=2)
            doc.add_paragraph(chapter["content"])

    # if there is a summary, add it to the document

    if "Samantekt" in response_json:
        doc.add_heading("Samantekt", level=2)
        doc.add_paragraph(response_json["Samantekt"])

    # Save the document to a temporary file using 'with'
    with NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        doc.save(temp_file.name)
        temp_file_path = temp_file.name  # Store the file path to return after closing

    return temp_file_path